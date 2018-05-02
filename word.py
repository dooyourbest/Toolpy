#!/usr/bin/env python
# coding=utf-8
from docx import Document
from docx.shared import Inches  
from docx.enum.style import WD_STYLE_TYPE
from docx import *
class Word:
	def __init__(self):
		self.document=Document()
	def addHeading(self,title,titleStyle=0):
		self.document.add_heading(title)
	def addParagraph(self,content):
		self.document.add_paragraph(content);
	def addPicture(self,path,width=Inches(6.25)):
		self.document.add_picture(path,width)
	def save(self,name):
		self.document.save(name)
	def addTable(self,arr,arrLimit=','):
		row=len(arr);
		col=len(arr[0].split(arrLimit))
		table=self.document.add_table(row,col,style='Light Grid Accent 3');
		for i in range(row):
		    hdr_cells=table.rows[i].cells
		    for j in range(col):
		    	newarr = arr[i].strip().split(arrLimit)
		    	hdr_cells[j].text=newarr[j]
    # def setStyle():

# word=Word()
# arr=['test,test1,test2','a,b,c','e,f,g']
# word.addHeading('123');
# word.addParagraph('helllo')
# word.addPicture('./picture/feature_27.png')
# word.addTable(arr)
# word.save('test1.docx')



  
#document = Document()  # 首先这是包的主要接口，这应该是利用的设计模式的一种，用来创建docx文档，里面也可以包含文档路径(d:\\2.docx)  
  
#document.add_heading('Document Title', 0)  # 这里是给文档添加一个标题，0表示 样式为title，1则为忽略，其他则是Heading{level},具体可以去<a href="https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html" target="_blank">官网</a>查;  
  
# p = document.add_paragraph('A plain paragraph having some ') # 这里是添加一个段落  
# p.add_run('bold').bold = True   # 这里是在这个段落p里文字some后面添加bold字符  
# p.add_run(' and some ')  
# p.add_run('italic.').italic = True  
  
# document.add_heading('Heading, level 1', level=1)   # 这里是添加标题1  
# document.add_paragraph('Intense quote', style='IntenseQuote') # 这里是添加段落，style后面则是样式  
  
# document.add_paragraph(  
#     'first item in unordered list', style='ListBullet'    # 添加段落，样式为unordered list类型  
# )  
# document.add_paragraph(  
#     'first item in ordered list', style='ListNumber'    # 添加段落，样式为ordered list数字类型  
# )  
  
# document.add_picture('monty-truth.png', width=Inches(1.25))  # 添加图片  
  
# table = document.add_table(rows=1, cols=3)   # 添加一个表格，每行三列  
# hdr_cells = table.rows[0].cells   # 表格第一行的所含有的所有列数  
# hdr_cells[0].text = 'Qty'    # 第一行的第一列,给这行里面添加文字  
# hdr_cells[1].text = 'Id'  
# hdr_cells[2].text = 'Desc'  
# for item in recordset:  
#     row_cells = table.add_row().cells   # 这是在这个表格第一行 (称作最后一行更好) 下面再添加新的一行  
#     row_cells[0].text = str(item.qty)  
#     row_cells[1].text = str(item.id)  
#     row_cells[2].text = item.desc  
   
# document.add_page_break()    # 添加分页符  
  
# document.save('demo.docx')  # 保存这个文档  